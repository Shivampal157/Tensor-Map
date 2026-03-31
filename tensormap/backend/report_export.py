"""Human-readable experiment reports (not model weights)."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any


def _safe(s: Any) -> str:
    if s is None:
        return ""
    t = str(s)
    return t.encode("utf-8", "replace").decode("utf-8", "replace")


def _for_pdf(s: Any) -> str:
    """FPDF core fonts are limited to latin-1."""
    return _safe(s).encode("latin-1", "replace").decode("latin-1")


def _node_lines(nodes: list[Any]) -> list[str]:
    lines: list[str] = []
    for n in nodes:
        if not isinstance(n, dict):
            lines.append(f"  • {_safe(n)}")
            continue
        data = n.get("data") or {}
        lt = data.get("layerType") or n.get("type") or "?"
        nid = n.get("id", "?")
        params = data.get("params") or {}
        try:
            pj = json.dumps(params, sort_keys=True, default=str)
        except TypeError:
            pj = str(params)
        lines.append(f"  • [{_safe(nid)}] {_safe(lt)}  params: {_safe(pj)}")
    return lines


def _edge_lines(edges: list[Any]) -> list[str]:
    lines: list[str] = []
    for e in edges:
        if not isinstance(e, dict):
            continue
        lines.append(f"  {_safe(e.get('source'))} -> {_safe(e.get('target'))}")
    return lines


def build_pdf_report(
    *,
    run_id: int,
    graph_name: str,
    nodes: list[Any],
    edges: list[Any],
    config: dict[str, Any],
    metrics_history: list[Any],
    created_at: datetime | None,
    out_path: str | Path,
) -> str:
    from fpdf import FPDF

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _for_pdf("TensorMap - training report"), ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 6, _for_pdf(f"Run ID: {run_id}    Graph: {graph_name}"), ln=True)
    if created_at:
        pdf.cell(0, 6, _for_pdf(f"Run created (UTC): {created_at.isoformat()}"), ln=True)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, _for_pdf("Model overview (canvas)"), ln=True)
    pdf.set_font("Helvetica", size=10)
    pdf.multi_cell(
        0,
        5,
        _for_pdf(
            "ONNX and TFLite are binary formats for apps and runtimes; "
            "this PDF is a human-readable summary."
        ),
        ln=1,
    )
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, _for_pdf("Nodes"), ln=True)
    pdf.set_font("Courier", size=8)
    for line in _node_lines(nodes):
        pdf.multi_cell(0, 4, _for_pdf(line), ln=1)
    pdf.ln(1)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, _for_pdf("Edges"), ln=True)
    pdf.set_font("Courier", size=8)
    for line in _edge_lines(edges):
        pdf.multi_cell(0, 4, _for_pdf(line), ln=1)
    pdf.ln(2)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, _for_pdf("Training configuration"), ln=True)
    pdf.set_font("Courier", size=8)
    pdf.multi_cell(
        0,
        4,
        _for_pdf(json.dumps(config, indent=2, sort_keys=True, default=str)),
        ln=1,
    )

    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, _for_pdf("Metrics by epoch"), ln=True)
    pdf.set_font("Courier", size=7)
    if not metrics_history:
        pdf.multi_cell(0, 4, _for_pdf("  (no rows)"), ln=1)
    else:
        hdr = metrics_history[0]
        if isinstance(hdr, dict):
            keys = sorted(hdr.keys())
            pdf.multi_cell(
                0,
                4,
                _for_pdf("  " + " | ".join(str(k) for k in keys)),
                ln=1,
            )
            for row in metrics_history:
                if not isinstance(row, dict):
                    continue
                vals = map(str, (_safe(row.get(k, "")) for k in keys))
                pdf.multi_cell(0, 4, _for_pdf("  " + " | ".join(vals)), ln=1)

    pdf.output(str(path))
    return str(path.resolve())


def build_docx_report(
    *,
    run_id: int,
    graph_name: str,
    nodes: list[Any],
    edges: list[Any],
    config: dict[str, Any],
    metrics_history: list[Any],
    created_at: datetime | None,
    out_path: str | Path,
) -> str:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    t = doc.add_heading("TensorMap – training report", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph(
        f"Run ID: {run_id}    Graph: {_safe(graph_name)}"
        + (f"    Created (UTC): {created_at.isoformat()}" if created_at else "")
    )
    doc.add_paragraph(
        "ONNX and TFLite are deployment formats (not viewable in Word). "
        "This document summarizes your experiment for sharing or archiving."
    )

    doc.add_heading("Nodes", level=1)
    for line in _node_lines(nodes):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Edges", level=1)
    for line in _edge_lines(edges):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Training configuration", level=1)
    doc.add_paragraph(_safe(json.dumps(config, indent=2, sort_keys=True, default=str)))

    doc.add_heading("Metrics by epoch", level=1)
    if not metrics_history or not isinstance(metrics_history[0], dict):
        doc.add_paragraph("(no metrics rows)")
    else:
        cols = sorted(metrics_history[0].keys())
        t2 = doc.add_table(rows=1 + len(metrics_history), cols=len(cols))
        t2.style = "Table Grid"
        for j, k in enumerate(cols):
            t2.rows[0].cells[j].text = _safe(k)
        for i, row in enumerate(metrics_history, start=1):
            if not isinstance(row, dict):
                continue
            for j, k in enumerate(cols):
                t2.rows[i].cells[j].text = _safe(row.get(k, ""))

    doc.save(str(path))
    return str(path.resolve())
